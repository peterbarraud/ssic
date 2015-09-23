#! /usr/bin/python
from docx import Document

doc = Document('teacher-list.docx')
target = open('table-out.txt', 'w')
rowindex = 0
for row in doc.tables[0].rows :
	if rowindex > 0:
		target.write('<tr>')
		cellindex = 0
		imgfilename = ''
		for  cell in row.cells :
			celltext = cell.text.strip()
			if cellindex == 1:
				imgfilename = cell.text.lower().replace(' ','')
			if cellindex > 0:
				if cellindex == 10 :
					target.write('<td><img height=100px src="images/' + imgfilename + '.png"></td>')
				else :
					if len(cell.paragraphs) == 1 or (len(cell.paragraphs) == 2 and cell.paragraphs[1].text.strip() == ''):
						if len(cell.paragraphs) == 1 :
							target.write('<td>' + cell.text + '</td>')
						else :
							target.write('<td>' + cell.paragraphs[0].text + '</td>')
					else :
						target.write('<td>')
						for para in cell.paragraphs :
							if para.text.strip() <> '' :
								target.write('<p>' + para.text + '</p>')
								#print cell.text
						target.write('</td>')
						
					
			cellindex = cellindex + 1
		target.write('</tr>')
		target.write("\n")	
	rowindex = rowindex + 1
target.close();

